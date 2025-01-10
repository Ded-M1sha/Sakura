from docx import Document


def find_topic(doc, topic_title):
    """
    Находит начало и конец темы по названию заголовка уровня 2.
    Возвращает список элементов, начиная с заголовка темы до следующего заголовка уровня 2 или уровня 1.
    """
    topic_content = []
    capture = False  # Флаг для захвата содержимого темы

    for paragraph in doc.paragraphs:
        # Поиск заголовка темы (уровень 2)
        if topic_title in paragraph.text and paragraph.style.name == 'Heading 2':
            capture = True  # Найден заголовок темы, начинаем захват
            topic_content.append(paragraph)
        elif capture and paragraph.style.name in ('Heading 1', 'Heading 2'):
            # Если начинается новый заголовок уровня 1 или 2, останавливаем захват
            break
        elif capture:
            # Добавляем параграфы, таблицы и изображения внутри темы
            topic_content.append(paragraph)

    return topic_content


def copy_topic_to_document(source_topic, target_doc):
    """
    Копирует элементы темы из одного документа в другой.
    """
    for element in source_topic:
        if isinstance(element, type(target_doc.add_paragraph())):
            # Копируем параграф
            p = target_doc.add_paragraph(element.text)
            p.style = element.style
        elif isinstance(element, type(target_doc.add_table(rows=1, cols=1))):
            # Копируем таблицу
            table = target_doc.add_table(rows=len(element.rows), cols=len(element.columns))
            for i, row in enumerate(element.rows):
                for j, cell in enumerate(row.cells):
                    table.cell(i, j).text = cell.text


def transfer_topic(source_path, target_path, topic_title):
    """
    Переносит тему из одного документа в другой на основе заголовка уровня 2.
    """
    # Открываем исходный и целевой документы
    source_doc = Document(source_path)
    target_doc = Document(target_path)

    # Находим нужную тему
    source_topic = find_topic(source_doc, topic_title)

    if source_topic:
        # Копируем тему в целевой документ
        copy_topic_to_document(source_topic, target_doc)

        # Сохраняем целевой документ
        target_doc.save("updated_" + target_path)
        print(f"Тема '{topic_title}' успешно перенесена в файл 'updated_{target_path}'")
    else:
        print(f"Тема с названием '{topic_title}' не найдена в документе '{source_path}'.")




# Пример использования
source_path = "Протокол_обследования_Онлайн_трейд_2_0.docx"
target_path = "шаблон_заполненный.docx"
chapter_title = "Общие характеристики номенклатуры"
transfer_topic(source_path, target_path, chapter_title)
